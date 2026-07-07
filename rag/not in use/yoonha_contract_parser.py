"""
yoonha_contract_parser.py
--------------------------
Workit — 정보시스템 개발구축 사업 표준계약서 파싱
PDF / DOCX 양식을 모두 지원한다.

출력 JSON 구조
--------------
{
  "contract_id"  : str,         # 파일명 stem
  "doc_type"     : "정보시스템_개발구축_표준계약서",
  "file_format"  : "pdf" | "docx",
  "전문": {
      "contract_name"         : str,
      "start_date"            : str,
      "end_date"              : str,
      "due_date"              : str,
      "contract_date"         : str,
      "total_amount"          : int | None,
      "supply_amount"         : int | None,
      "vat"                   : int | None,
      "payment_schedule"      : [
          {"type": str, "rate": str,
           "amount": int|None, "date": str, "method": str}, ...
      ],
      "performance_bond"      : int | None,
      "performance_bond_rate" : float | None,
      "task_confirm_date"     : str,
      "delay_rate"            : float | None,
      "penalty_limit"         : int | None,
      "penalty_limit_rate"    : float | None,
      "warranty_bond"         : int | None,
      "warranty_bond_rate"    : float | None,
      "warranty_period"       : str,
      "client": {"name", "ceo", "addr", "tel", "biz_no"},
      "vendor": {"name", "ceo", "addr", "tel", "biz_no"},
  },
  "본문": [
      {"article_id", "article_number", "title", "text"}, ...
  ],
  "raw_issues": [str, ...]   # ※[검토필요] 배너 (문제 계약서)
}

실행:
    # 단일 파일
    python yoonha_contract_parser.py path/to/contract.pdf

    # 폴더 일괄
    python yoonha_contract_parser.py path/to/folder/
"""

from __future__ import annotations
import re, sys, json
from pathlib import Path

# ══════════════════════════════════════════════
# 유틸
# ══════════════════════════════════════════════

def _parse_amount(text: str) -> int | None:
    t = re.sub(r"[,원정￦\s금]", "", text)
    m = re.search(r"\d+", t)
    return int(m.group()) if m else None

def _parse_rate(text: str) -> float | None:
    m = re.search(r"(\d+(?:\.\d+)?)\s*%", text)
    return float(m.group(1)) if m else None

def _parse_date(text: str) -> str:
    m = re.search(r"\d{4}\.\d{2}\.\d{2}", text)
    return m.group() if m else text.strip()

# 키 정규화 맵
KEY_MAP = {
    "1. 계약명"              : "계약명",
    "2. 계약기간"            : "계약기간",
    "3. 납기일"              : "납기일",
    "4. 계약금액"            : "계약금액",
    "공급가액"               : "공급가액",
    "부가가치세"             : "부가가치세",
    "5. 계약이행보증금"      : "계약이행보증금",
    "6. 과업내용서 발급시기" : "과업내용서발급시기",
    "7. 지연이자요율"        : "지연이자요율",
    "8. 지체상금요율"        : "지체상금요율",
    "9. 지체상금 한도"       : "지체상금한도",
    "12. 하자보수보증금"     : "하자보수보증금",
    "하자담보책임기간"       : "하자담보책임기간",
}

PAYMENT_TYPES = {
    "선급금","중도금","잔금","합계",
    "중도금 1차","중도금 2차","중도금 3차",
}

PARTY_PREFIXES = ("상호:","대표자:","주소:","전화:","사업자번호:")

_ARTICLE_PAT = re.compile(r"^(제\d+조)\(([^)]+)\)\s*(.*)", re.DOTALL)


# ══════════════════════════════════════════════
# 테이블 행 → kv / payment / party 분류
# ══════════════════════════════════════════════

def _classify_tables(tables: list[list[list[str]]]) -> tuple[
    dict[str,str], list[list[str]], dict[str,str], dict[str,str]
]:
    kv: dict[str,str] = {}
    payment_rows: list[list[str]] = []
    client_kv: dict[str,str] = {}
    vendor_kv: dict[str,str] = {}

    for table in tables:
        if not table:
            continue

        # 당사자 테이블
        if table[0] == ["발주자", "공급자"]:
            for row in table[1:]:
                if len(row) < 2:
                    continue
                left, right = row[0].strip(), row[1].strip()
                for pf in PARTY_PREFIXES:
                    if left.startswith(pf):
                        client_kv[pf.rstrip(":")] = left[len(pf):].strip()
                    if right.startswith(pf):
                        vendor_kv[pf.rstrip(":")] = right[len(pf):].strip()
            continue

        for row in table:
            if not row or not row[0].strip():
                continue
            k0 = row[0].strip()
            v0 = row[1].strip() if len(row) > 1 else ""

            # 대금 지급 행
            if k0.rstrip("*") in PAYMENT_TYPES:
                payment_rows.append([c.strip() for c in row])
                continue

            # 전문 키-값
            for raw_key, norm_key in KEY_MAP.items():
                if k0.startswith(raw_key):
                    kv[norm_key] = v0
                    break

    return kv, payment_rows, client_kv, vendor_kv


# ══════════════════════════════════════════════
# 전문 구조 조립
# ══════════════════════════════════════════════

def _build_전문(
    kv: dict[str,str],
    payment_rows: list[list[str]],
    client_kv: dict[str,str],
    vendor_kv: dict[str,str],
    contract_date: str,
) -> dict:

    # 기간
    period = kv.get("계약기간","")
    dates  = re.findall(r"\d{4}\.\d{2}\.\d{2}", period)
    start_date = dates[0] if len(dates)>=1 else ""
    end_date   = dates[1] if len(dates)>=2 else ""

    # 대금 지급 계획
    payment_schedule = []
    for row in payment_rows:
        ptype = row[0] if len(row)>0 else ""
        if ptype in ("구분","합계"):
            continue
        rate  = row[1] if len(row)>1 else ""
        amt   = _parse_amount(row[2]) if len(row)>2 else None
        date  = row[3] if len(row)>3 else ""
        meth  = row[4] if len(row)>4 else ""
        payment_schedule.append({
            "type":ptype, "rate":rate,
            "amount":amt, "date":date, "method":meth
        })

    # 보증금 / 한도
    pb_raw = kv.get("계약이행보증금","")
    pl_raw = kv.get("지체상금한도","")
    wb_raw = kv.get("하자보수보증금","")

    def _extract_party(pkv: dict[str,str]) -> dict:
        return {
            "name"  : pkv.get("상호",""),
            "ceo"   : re.sub(r"\s*\(인\)","", pkv.get("대표자","")).strip(),
            "addr"  : pkv.get("주소",""),
            "tel"   : pkv.get("전화",""),
            "biz_no": pkv.get("사업자번호",""),
        }

    return {
        "contract_name"        : kv.get("계약명",""),
        "start_date"           : start_date,
        "end_date"             : end_date,
        "due_date"             : _parse_date(kv.get("납기일","")),
        "contract_date"        : contract_date,
        "total_amount"         : _parse_amount(kv.get("계약금액","")),
        "supply_amount"        : _parse_amount(kv.get("공급가액","")),
        "vat"                  : _parse_amount(kv.get("부가가치세","")),
        "payment_schedule"     : payment_schedule,
        "performance_bond"     : _parse_amount(pb_raw),
        "performance_bond_rate": _parse_rate(pb_raw),
        "task_confirm_date"    : _parse_date(kv.get("과업내용서발급시기","")),
        "delay_rate"           : _parse_rate(kv.get("지연이자요율","")),
        "penalty_limit"        : _parse_amount(pl_raw),
        "penalty_limit_rate"   : _parse_rate(pl_raw),
        "warranty_bond"        : _parse_amount(wb_raw),
        "warranty_bond_rate"   : _parse_rate(wb_raw),
        "warranty_period"      : kv.get("하자담보책임기간",""),
        "client"               : _extract_party(client_kv),
        "vendor"               : _extract_party(vendor_kv),
    }


# ══════════════════════════════════════════════
# 본문 조항 파싱
# ══════════════════════════════════════════════

def _parse_articles(lines: list[str]) -> list[dict]:
    articles = []
    cur_id = cur_title = None
    cur_buf: list[str] = []

    def flush():
        if cur_id:
            articles.append({
                "article_id"    : cur_id,
                "article_number": cur_id,
                "title"         : cur_title,
                "text"          : " ".join(cur_buf).strip(),
            })

    for line in lines:
        line = line.strip()
        if not line:
            continue
        m = _ARTICLE_PAT.match(line)
        if m:
            flush()
            cur_id    = m.group(1)
            cur_title = m.group(2)
            cur_buf   = [m.group(3)] if m.group(3) else []
        elif cur_id:
            cur_buf.append(line)

    flush()
    return articles


# ══════════════════════════════════════════════
# PDF 파서
# ══════════════════════════════════════════════

def _parse_pdf(path: Path) -> dict:
    try:
        import pdfplumber
    except ImportError:
        raise ImportError("pdfplumber 미설치: pip install pdfplumber")

    all_tables: list[list[list[str]]] = []
    all_text_lines: list[str] = []

    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            # 표 추출
            for tbl in page.extract_tables() or []:
                cleaned = [
                    [c.strip() if c else "" for c in row]
                    for row in tbl if any(c and c.strip() for c in row)
                ]
                if cleaned:
                    all_tables.append(cleaned)
            # 텍스트 라인 (본문 조항 + 이슈 배너 + 계약일용)
            text = page.extract_text() or ""
            all_text_lines.extend(text.splitlines())

    # 이슈 배너
    raw_issues = [
        re.sub(r"^※\s*\[검토필요\]\s*","", l).strip()
        for l in all_text_lines if "검토필요" in l
    ]

    # 계약 체결일
    contract_date = ""
    for l in all_text_lines:
        if l.strip().startswith("계약 체결일:"):
            contract_date = _parse_date(l.replace("계약 체결일:",""))
            break

    # 본문 조항
    article_lines = []
    collecting = False
    for l in all_text_lines:
        ls = l.strip()
        if _ARTICLE_PAT.match(ls):
            collecting = True
        if collecting and ls:
            article_lines.append(ls)

    # 테이블 → kv
    kv, payment_rows, client_kv, vendor_kv = _classify_tables(all_tables)

    전문 = _build_전문(kv, payment_rows, client_kv, vendor_kv, contract_date)
    본문 = _parse_articles(article_lines)

    return {
        "contract_id": path.stem,
        "doc_type"   : "정보시스템_개발구축_표준계약서",
        "file_format": "pdf",
        "전문"       : 전문,
        "본문"       : 본문,
        "raw_issues" : raw_issues,
    }


# ══════════════════════════════════════════════
# DOCX 파서
# ══════════════════════════════════════════════

def _parse_docx(path: Path) -> dict:
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx 미설치: pip install python-docx")

    doc = Document(str(path))

    raw_issues: list[str] = []
    article_lines: list[str] = []
    contract_date = ""
    collecting_articles = False

    for para in doc.paragraphs:
        t = para.text.strip()
        if not t:
            continue
        if "검토필요" in t:
            raw_issues.append(re.sub(r"^※\s*\[검토필요\]\s*","",t).strip())
            continue
        if t.startswith("계약 체결일:"):
            contract_date = _parse_date(t.replace("계약 체결일:",""))
            continue
        if _ARTICLE_PAT.match(t):
            collecting_articles = True
        if collecting_articles:
            article_lines.append(t)

    # 테이블 → kv
    all_tables: list[list[list[str]]] = []
    for table in doc.tables:
        rows = [[c.text.strip() for c in row.cells] for row in table.rows]
        all_tables.append(rows)

    kv, payment_rows, client_kv, vendor_kv = _classify_tables(all_tables)

    전문 = _build_전문(kv, payment_rows, client_kv, vendor_kv, contract_date)
    본문 = _parse_articles(article_lines)

    return {
        "contract_id": path.stem,
        "doc_type"   : "정보시스템_개발구축_표준계약서",
        "file_format": "docx",
        "전문"       : 전문,
        "본문"       : 본문,
        "raw_issues" : raw_issues,
    }


# ══════════════════════════════════════════════
# 공개 API
# ══════════════════════════════════════════════

def parse_contract(file_path: str | Path) -> dict:
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"파일 없음: {path}")
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _parse_pdf(path)
    elif suffix in (".docx", ".doc"):
        return _parse_docx(path)
    else:
        raise ValueError(f"지원하지 않는 형식: {suffix}")


def save_parsed_contract(result: dict, out_path: str | Path) -> Path:
    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    with open(out_path, "w", encoding="utf-8") as f:
        json.dump(result, f, ensure_ascii=False, indent=2)
    return out_path


def parse_contract_dir(
    input_dir: str | Path,
    output_dir: str | Path,
) -> dict[str, Path]:
    input_dir  = Path(input_dir)
    output_dir = Path(output_dir)

    files = sorted([
        f for f in input_dir.iterdir()
        if f.suffix.lower() in (".pdf", ".docx", ".doc")
    ])
    if not files:
        print(f"[contract_parser] 처리할 파일 없음: {input_dir}")
        return {}

    print(f"[contract_parser] {len(files)}개 파일 파싱 시작")
    print("=" * 60)
    results = {}
    for fp in files:
        try:
            print(f"  파싱 중: {fp.name} ...", end=" ", flush=True)
            parsed   = parse_contract(fp)
            out_path = output_dir / (fp.stem + ".json")
            save_parsed_contract(parsed, out_path)
            n_art = len(parsed["본문"])
            n_iss = len(parsed["raw_issues"])
            print(f"본문 {n_art}조 | 이슈 {n_iss}건 → {out_path.name}")
            results[fp.name] = out_path
        except Exception as e:
            print(f"실패: {e}")

    print(f"\n[contract_parser] 완료: {len(results)}/{len(files)}개 처리")
    print(f"[contract_parser] 저장 위치: {output_dir}")
    return results


# ══════════════════════════════════════════════
# CLI
# ══════════════════════════════════════════════

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print(__doc__)
        sys.exit(0)

    target = Path(sys.argv[1])

    if target.is_dir():
        out_dir = target.parent / (target.name + "_parsed")
        parse_contract_dir(target, out_dir)
    elif target.is_file():
        parsed   = parse_contract(target)
        out_path = target.parent / (target.stem + ".json")
        save_parsed_contract(parsed, out_path)
        print(f"저장 완료: {out_path}")
        print(json.dumps(parsed, ensure_ascii=False, indent=2))
    else:
        print(f"[오류] 파일 또는 폴더를 찾을 수 없습니다: {target}")
        sys.exit(1)